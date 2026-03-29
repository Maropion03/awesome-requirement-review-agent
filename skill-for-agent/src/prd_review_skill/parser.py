"""PRD Parser Module.

Parses PRD documents (Markdown/DOCX) and extracts structured information.
"""

import re
from dataclasses import dataclass, field
from typing import List, Dict, Any, Optional
from pathlib import Path

from .models import PRDDocument, PRDSection


class PRDParser:
    """Parser for PRD documents."""

    def parse(self, file_path: str) -> PRDDocument:
        """Parse a PRD document from file path.

        Args:
            file_path: Path to the PRD file (.md or .docx)

        Returns:
            PRDDocument object containing structured information
        """
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix == ".md":
            return self._parse_markdown(path)
        elif suffix == ".docx":
            return self._parse_docx(path)
        else:
            raise ValueError(f"Unsupported file format: {suffix}")

    def _parse_markdown(self, path: Path) -> PRDDocument:
        """Parse a Markdown PRD file."""
        with open(path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.split("\n")
        metadata = self._extract_metadata(lines)
        sections = self._extract_sections(lines)

        return PRDDocument(
            project_name=metadata["project_name"],
            version=metadata["version"],
            date=metadata["date"],
            author=metadata["author"],
            sections=sections,
            raw_text=content,
            file_format="markdown",
        )

    def _parse_docx(self, path: Path) -> PRDDocument:
        """Parse a DOCX PRD file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. Install with: pip install python-docx"
            )

        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        content = "\n".join(paragraphs)
        lines = content.split("\n")

        metadata = self._extract_metadata(lines)
        sections = self._extract_sections(lines)

        return PRDDocument(
            project_name=metadata["project_name"],
            version=metadata["version"],
            date=metadata["date"],
            author=metadata["author"],
            sections=sections,
            raw_text=content,
            file_format="docx",
        )

    def _extract_metadata(self, lines: List[str]) -> Dict[str, Optional[str]]:
        """Extract metadata from the document header."""
        metadata = {
            "project_name": "未知项目",
            "version": "1.0",
            "date": None,
            "author": None,
        }

        for i, line in enumerate(lines[:20]):
            line = line.strip()

            # Look for project name
            name_match = re.search(r"项目[：:]\s*(.+)", line)
            if name_match:
                metadata["project_name"] = name_match.group(1).strip()

            # Look for version
            version_match = re.search(r"版本[：:]\s*(v?\d+[\d.]*)", line, re.IGNORECASE)
            if version_match:
                metadata["version"] = version_match.group(1).strip()

            # Look for date
            date_match = re.search(r"日期[：:]\s*(\d{4}[-/]\d{2}[-/]\d{2})", line)
            if date_match:
                metadata["date"] = date_match.group(1).strip()

            # Look for author
            author_match = re.search(r"作者[：:]\s*(.+)", line)
            if author_match:
                metadata["author"] = author_match.group(1).strip()

        return metadata

    def _extract_sections(self, lines: List[str]) -> List[PRDSection]:
        """Extract sections from markdown lines."""
        sections = []
        current_section = None
        current_content_lines = []
        current_level = 0

        for line_num, line in enumerate(lines, 1):
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())

            if heading_match:
                if current_section:
                    current_section.content = "\n".join(current_content_lines).strip()
                    current_section.children = self._build_section_tree(
                        current_content_lines, current_level + 1
                    )

                level = len(heading_match.group(1))
                title = heading_match.group(2).strip()

                current_section = PRDSection(
                    title=title,
                    level=level,
                    content="",
                    line_number=line_num,
                )
                sections.append(current_section)
                current_content_lines = []
                current_level = level
            else:
                current_content_lines.append(line)

        if current_section:
            current_section.content = "\n".join(current_content_lines).strip()
            current_section.children = self._build_section_tree(
                current_content_lines, current_level + 1
            )

        return sections

    def _build_section_tree(
        self, lines: List[str], parent_level: int
    ) -> List[PRDSection]:
        """Build a tree of subsections."""
        sections = []
        current_section = None
        current_content_lines = []

        for line in lines:
            heading_match = re.match(r"^(#{1,6})\s+(.+)$", line.strip())

            if heading_match:
                level = len(heading_match.group(1))

                if level <= parent_level:
                    break

                if current_section:
                    current_section.content = "\n".join(current_content_lines).strip()
                    sections.append(current_section)

                title = heading_match.group(2).strip()
                current_section = PRDSection(
                    title=title,
                    level=level,
                    content="",
                    line_number=0,
                )
                current_content_lines = []
            else:
                if current_section is None:
                    pass
                else:
                    current_content_lines.append(line)

        if current_section:
            current_section.content = "\n".join(current_content_lines).strip()
            sections.append(current_section)

        return sections

    def extract_key_info(self, doc: PRDDocument) -> Dict[str, Any]:
        """Extract key information from parsed PRD."""
        info = {
            "sections_count": len(doc.sections),
            "total_words": len(doc.raw_text),
            "has_acceptance_criteria": False,
            "priorities_found": [],
            "technical_deps": [],
        }

        raw = doc.raw_text.lower()

        if re.search(r"验收|acceptance|criteria|验收标准", raw):
            info["has_acceptance_criteria"] = True

        priority_matches = re.findall(r"P[0-3]|优先级[：:]\s*[0-3]", doc.raw_text)
        info["priorities_found"] = list(set(priority_matches))

        dep_keywords = ["依赖", "接口", "API", "service", "第三方", "系统", "platform"]
        for keyword in dep_keywords:
            if keyword.lower() in raw:
                info["technical_deps"].append(keyword)

        return info


def parse_prd(file_path: str) -> PRDDocument:
    """Convenience function to parse a PRD file."""
    parser = PRDParser()
    return parser.parse(file_path)
