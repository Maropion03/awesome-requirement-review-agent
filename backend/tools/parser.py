"""
PRD Parser - parses Markdown and DOCX files.
"""

import re
from pathlib import Path
from typing import Optional, Dict, Any


class PRDParser:
    """Parser for PRD documents in Markdown and DOCX formats."""

    def parse(self, file_path: str) -> str:
        """
        Parse a PRD file and return its content.

        Args:
            file_path: Path to the PRD file (.md or .docx)

        Returns:
            Extracted text content
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        suffix = path.suffix.lower()

        if suffix == ".md":
            return self.parse_markdown(path)
        elif suffix == ".docx":
            return self.parse_docx(path)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")

    def parse_markdown(self, path: Path) -> str:
        """Parse a Markdown PRD file."""
        content = path.read_text(encoding="utf-8")
        return self._clean_markdown(content)

    def parse_docx(self, path: Path) -> str:
        """Parse a DOCX PRD file."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is required for DOCX parsing")

        doc = Document(path)
        paragraphs = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    paragraphs.append(" | ".join(row_text))

        return self._clean_markdown("\n".join(paragraphs))

    def _clean_markdown(self, content: str) -> str:
        """Clean markdown content for LLM processing."""
        # Remove markdown image syntax
        content = re.sub(r"!\[.*?\]\(.*?\)", "", content)

        # Remove markdown link syntax but keep text
        content = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", content)

        # Remove HTML tags
        content = re.sub(r"<[^>]+>", "", content)

        # Remove code blocks (keep content for context)
        content = re.sub(r"```[\s\S]*?```", "", content)

        # Remove inline code
        content = re.sub(r"`[^`]+`", "", content)

        # Normalize whitespace
        content = re.sub(r"\n{3,}", "\n\n", content)
        content = re.sub(r" {2,}", " ", content)

        return content.strip()

    def extract_metadata(self, content: str) -> Dict[str, Any]:
        """
        Extract metadata from PRD content.

        Args:
            content: PRD text content

        Returns:
            Dictionary of extracted metadata
        """
        metadata = {
            "title": self._extract_title(content),
            "version": self._extract_version(content),
            "author": self._extract_author(content),
            "date": self._extract_date(content),
        }
        return metadata

    def _extract_title(self, content: str) -> Optional[str]:
        """Extract project title."""
        # Check first heading
        match = re.search(r"^#\s+(.+)$", content, re.MULTILINE)
        if match:
            return match.group(1).strip()

        # Check title in first line
        lines = content.split("\n")
        if lines:
            return lines[0].strip()[:100]

        return None

    def _extract_version(self, content: str) -> Optional[str]:
        """Extract version information."""
        patterns = [
            r"版本[：:]\s*v?(\d+\.\d+(?:\.\d+)?)",
            r"version[：:]\s*v?(\d+\.\d+(?:\.\d+)?)",
            r"V(\d+\.\d+(?:\.\d+)?)",
        ]

        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1)

        return None

    def _extract_author(self, content: str) -> Optional[str]:
        """Extract author information."""
        patterns = [
            r"作者[：:]\s*(.+)",
            r"author[：:]\s*(.+)",
            r"编写人[：:]\s*(.+)",
        ]

        for pattern in patterns:
            match = re.search(pattern, content, re.IGNORECASE)
            if match:
                return match.group(1).strip()

        return None

    def _extract_date(self, content: str) -> Optional[str]:
        """Extract date information."""
        patterns = [
            r"日期[：:]\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
            r"date[：:]\s*(\d{4}[-/]\d{1,2}[-/]\d{1,2})",
            r"(\d{4}年\d{1,2}月\d{1,2}日)",
        ]

        for pattern in patterns:
            match = re.search(pattern, content)
            if match:
                return match.group(1)

        return None

    def extract_sections(self, content: str) -> Dict[str, str]:
        """
        Extract sections from PRD content.

        Args:
            content: PRD text content

        Returns:
            Dictionary mapping section names to content
        """
        sections = {}

        # Split by major headings (## and ###)
        parts = re.split(r"^#{2,3}\s+(.+)$", content, flags=re.MULTILINE)

        if len(parts) > 1:
            # parts[0] is content before first heading
            current_section = "_preamble"
            sections[current_section] = parts[0].strip()

            # Process heading-content pairs
            for i in range(1, len(parts), 2):
                if i + 1 < len(parts):
                    heading = parts[i].strip()
                    section_content = parts[i + 1].strip()
                    sections[heading] = section_content

        return sections
